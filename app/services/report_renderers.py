"""Report renderers: pure functions from the canonical model to each format.

Canary: Plainspoken Marmot

Every renderer takes a `report_payload` dict (the canonical, schema-governed model
from `report_builder`) and emits one format. JSON is the model verbatim; MD / CSV /
DOCX are derived from it. Charts are stored as data series in the model and drawn
to PNG here (matplotlib, headless Agg backend); the series are also emitted as
tables so CSV/MD stay complete.

Hard rule: downloaded artifacts (CSV/DOCX/MD) render the COMPLETE dataset — every
row, sorted — with no top-N truncation. Only `render_html` (the on-screen preview)
may elide.
"""

from __future__ import annotations

import csv
import html
import io
import json
from typing import Any, Dict, List, Optional, Tuple

import matplotlib

matplotlib.use("Agg")  # headless: no display on the server
import matplotlib.pyplot as plt  # noqa: E402

_TABULAR = {"table", "ranked_list", "chart"}


# --------------------------------------------------------------------------- #
# Canonical-model helpers: reduce a section to (header, rows) for tabular forms #
# --------------------------------------------------------------------------- #
def section_to_table(section: Dict[str, Any]) -> Optional[Tuple[List[str], List[List[Any]]]]:
    """Return (header, rows) for a tabular section, or None if not tabular.

    `ranked_list` and `chart` (its data series) are flattened to tables so the
    complete data appears in CSV/MD, not just `table` sections.
    """
    t = section.get("type")
    body = section.get("body") or {}
    if t == "table":
        cols = body.get("columns") or []
        header = [c["label"] for c in cols]
        keys = [c["key"] for c in cols]
        rows = [[r.get(k) for k in keys] for r in body.get("rows") or []]
        return header, rows
    if t == "ranked_list":
        header = ["Rank", "Item", "Median", "IQR", "Agreement"]
        rows = []
        for it in body.get("items") or []:
            s = it.get("stats") or {}
            rows.append(
                [it.get("rank"), it.get("label"), s.get("median"), s.get("iqr"),
                 s.get("agreement_band")]
            )
        return header, rows
    if t == "chart":
        cats = (body.get("x_axis") or {}).get("categories") or []
        header = ["Series"] + [str(c) for c in cats]
        rows = [[s.get("name")] + list(s.get("points") or []) for s in body.get("series") or []]
        return header, rows
    return None


# --------------------------------------------------------------------------- #
# JSON                                                                         #
# --------------------------------------------------------------------------- #
def render_json(report: Dict[str, Any]) -> str:
    """Canonical model, verbatim."""
    return json.dumps(report, indent=2, ensure_ascii=False)


# --------------------------------------------------------------------------- #
# CSV (one table per tabular section; complete rows)                           #
# --------------------------------------------------------------------------- #
def render_csv_files(report: Dict[str, Any]) -> Dict[str, str]:
    """Return {section_id: csv_text} for every tabular section, no truncation."""
    out: Dict[str, str] = {}
    for sec in report.get("sections") or []:
        tbl = section_to_table(sec)
        if tbl is None:
            continue
        header, rows = tbl
        buf = io.StringIO()
        w = csv.writer(buf)
        w.writerow(header)
        for r in rows:
            w.writerow(["" if v is None else v for v in r])
        out[sec["id"]] = buf.getvalue()
    return out


def render_csv(report: Dict[str, Any], section_id: Optional[str] = None) -> str:
    """One CSV: the named tabular section, or the primary trajectory/table by default."""
    files = render_csv_files(report)
    if section_id:
        if section_id not in files:
            raise KeyError(f"no tabular section {section_id!r}")
        return files[section_id]
    for pref in ("trajectory", "final_ranking"):
        if pref in files:
            return files[pref]
    return next(iter(files.values())) if files else ""


# --------------------------------------------------------------------------- #
# Markdown (complete tables; charts noted + series tabulated)                  #
# --------------------------------------------------------------------------- #
def _md_table(header: List[str], rows: List[List[Any]]) -> List[str]:
    out = ["| " + " | ".join(str(h) for h in header) + " |",
           "| " + " | ".join("---" for _ in header) + " |"]
    for r in rows:
        out.append("| " + " | ".join("" if v is None else str(v) for v in r) + " |")
    return out


def render_markdown(report: Dict[str, Any]) -> str:
    lines: List[str] = [f"# {report.get('title', 'Meeting Report')}", ""]
    m = report.get("meeting") or {}
    if m.get("method"):
        lines.append(f"_Method: {m['method'].get('name')} {m['method'].get('version') or ''}_".rstrip())
    lines.append("")
    for sec in report.get("sections") or []:
        lines.append(f"## {sec.get('title')}")
        t, body = sec.get("type"), sec.get("body") or {}
        if t == "narrative":
            lines.append(body.get("markdown", ""))
        elif t == "key_value":
            lines += [f"- **{p['label']}**: {p['value']}" for p in body.get("pairs") or []]
        elif t in ("table", "ranked_list", "chart"):
            if t == "chart":
                lines.append(f"_Chart ({body.get('chart_kind')}); series below._")
            header, rows = section_to_table(sec)
            lines += _md_table(header, rows)
        elif t == "comment_thread":
            for g in body.get("groups") or []:
                lines.append(f"**{g.get('heading')}**")
                lines += [f"- {c.get('text')}" for c in g.get("comments") or []]
        elif t == "rounds":
            for rd in body.get("rounds") or []:
                lines.append(f"- **Round {rd.get('round_number')}**: {rd.get('summary', '')}")
        lines.append("")
    return "\n".join(lines)


# --------------------------------------------------------------------------- #
# HTML preview (may elide long tables; downloads remain complete)              #
# --------------------------------------------------------------------------- #
def render_html(report: Dict[str, Any], *, max_rows: int = 12) -> str:
    """Render an on-screen preview of the canonical model.

    Unlike downloaded artifacts, this preview may elide long tabular sections so
    the meeting page stays scannable. It still reads only the stored model.
    """
    parts: List[str] = [
        '<article class="report-preview-document">',
        f"<h2>{html.escape(str(report.get('title') or 'Meeting Report'))}</h2>",
    ]
    meeting = report.get("meeting") or {}
    method = meeting.get("method") or {}
    meta_bits = []
    if method.get("name"):
        meta_bits.append(str(method.get("name")))
    if meeting.get("round_count") is not None:
        meta_bits.append(f"{meeting.get('round_count')} rounds")
    if meta_bits:
        parts.append(f'<p class="report-preview-meta">{html.escape(" | ".join(meta_bits))}</p>')

    for sec in report.get("sections") or []:
        title = html.escape(str(sec.get("title") or "Section"))
        section_type = sec.get("type")
        body = sec.get("body") or {}
        parts.append(f'<section class="report-preview-section" data-section-type="{html.escape(str(section_type or ""))}">')
        parts.append(f"<h3>{title}</h3>")
        if section_type == "narrative":
            text = str(body.get("markdown") or "")
            for paragraph in [p.strip() for p in text.split("\n\n") if p.strip()]:
                parts.append(f"<p>{html.escape(paragraph)}</p>")
        elif section_type == "key_value":
            parts.append("<dl>")
            for pair in body.get("pairs") or []:
                parts.append(f"<dt>{html.escape(str(pair.get('label') or ''))}</dt>")
                parts.append(f"<dd>{html.escape(str(pair.get('value') or ''))}</dd>")
            parts.append("</dl>")
        elif section_type in _TABULAR:
            if section_type == "chart":
                parts.append("<p>Chart preview data:</p>")
            table = section_to_table(sec)
            if table is not None:
                header, rows = table
                shown = rows[:max_rows]
                parts.append("<div class=\"report-preview-table-wrap\"><table>")
                parts.append(
                    "<thead><tr>"
                    + "".join(f"<th>{html.escape(str(h))}</th>" for h in header)
                    + "</tr></thead>"
                )
                parts.append("<tbody>")
                for row in shown:
                    parts.append(
                        "<tr>"
                        + "".join(
                            f"<td>{'' if value is None else html.escape(str(value))}</td>"
                            for value in row
                        )
                        + "</tr>"
                    )
                if len(rows) > len(shown):
                    parts.append(
                        f'<tr class="report-preview-elided"><td colspan="{len(header)}">'
                        f"Showing {len(shown)} of {len(rows)} rows. Download for the complete table."
                        "</td></tr>"
                    )
                parts.append("</tbody></table></div>")
        elif section_type == "comment_thread":
            for group in body.get("groups") or []:
                parts.append(f"<h4>{html.escape(str(group.get('heading') or 'Comments'))}</h4>")
                parts.append("<ul>")
                for comment in group.get("comments") or []:
                    parts.append(f"<li>{html.escape(str(comment.get('text') or ''))}</li>")
                parts.append("</ul>")
        elif section_type == "rounds":
            parts.append("<ol>")
            for round_info in body.get("rounds") or []:
                summary = html.escape(str(round_info.get("summary") or ""))
                number = html.escape(str(round_info.get("round_number") or ""))
                parts.append(f"<li><strong>Round {number}</strong>: {summary}</li>")
            parts.append("</ol>")
        parts.append("</section>")
    parts.append("</article>")
    return "\n".join(parts)


# --------------------------------------------------------------------------- #
# Charts (data series -> PNG bytes)                                            #
# --------------------------------------------------------------------------- #
def render_chart_png(section: Dict[str, Any]) -> bytes:
    """Render a `chart` section's series to a PNG (headless)."""
    body = section.get("body") or {}
    kind = body.get("chart_kind")
    cats = (body.get("x_axis") or {}).get("categories") or []
    series = body.get("series") or []
    fig, ax = plt.subplots(figsize=(5, 3))
    try:
        if kind == "stacked_bar":
            bottom = [0.0] * len(cats)
            for s in series:
                pts = [p or 0 for p in s.get("points") or []]
                ax.bar(cats, pts, bottom=bottom, label=s.get("name"))
                bottom = [b + p for b, p in zip(bottom, pts)]
            ax.legend(fontsize=7)
        else:  # line / bump
            for s in series:
                ax.plot(cats, s.get("points"), marker="o", label=s.get("name"))
            if kind == "bump":
                ax.invert_yaxis()
            else:
                ax.legend(fontsize=7)
        ax.set_title(section.get("title", ""))
        ya = body.get("y_axis") or {}
        if ya.get("label"):
            ax.set_ylabel(ya["label"])
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=90)
        return buf.getvalue()
    finally:
        plt.close(fig)


# --------------------------------------------------------------------------- #
# DOCX (complete tables + embedded chart PNGs)                                 #
# --------------------------------------------------------------------------- #
def render_docx(report: Dict[str, Any]) -> bytes:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading(report.get("title", "Meeting Report"), level=0)
    for sec in report.get("sections") or []:
        doc.add_heading(sec.get("title", ""), level=1)
        t, body = sec.get("type"), sec.get("body") or {}
        if t == "narrative":
            doc.add_paragraph(body.get("markdown", ""))
        elif t == "key_value":
            for p in body.get("pairs") or []:
                doc.add_paragraph(f"{p['label']}: {p['value']}", style="List Bullet")
        elif t in ("table", "ranked_list"):
            header, rows = section_to_table(sec)
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for i, h in enumerate(header):
                table.rows[0].cells[i].text = str(h)
            for r in rows:
                cells = table.add_row().cells
                for i, v in enumerate(r):
                    cells[i].text = "" if v is None else str(v)
        elif t == "chart":
            doc.add_picture(io.BytesIO(render_chart_png(sec)), width=Inches(4.5))
            header, rows = section_to_table(sec)  # series table beside the image
            table = doc.add_table(rows=1, cols=len(header))
            for i, h in enumerate(header):
                table.rows[0].cells[i].text = str(h)
            for r in rows:
                cells = table.add_row().cells
                for i, v in enumerate(r):
                    cells[i].text = "" if v is None else str(v)
        elif t == "rounds":
            for rd in body.get("rounds") or []:
                doc.add_paragraph(
                    f"Round {rd.get('round_number')}: {rd.get('summary', '')}",
                    style="List Bullet",
                )
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
